import fitz
import os
import hashlib
import requests
import chromadb
import docx
import subprocess
import tempfile
import shutil

from config import (
    PASTA_DOCUMENTOS,
    TAMANHO_CHUNK,
    SOBREPOSICAO,
    PASTA_BANCO_VETORIAL,
    VERSAO_PIPELINE,
    TAMANHO_LOTE_EMBEDDING,
    MODELO_EMBEDDING,
)

def tabela_para_prosa(tabela) -> str:
    """Converte uma tabela achada por `pagina.find_tables()` em frases "Linha N: cabeçalho=valor"
    — formato genérico (não presume o significado da tabela), gerado pra QUALQUER tabela que o
    PyMuPDF detecte, não só a que motivou isso.

    Achado real, 2026-08-25 (datasheet de CI, tabela-verdade binária do CD4051B): `get_text()`
    achata tabela numa sequência linear de números, sem coluna nem linha — o modelo lia errado de
    forma consistente mesmo com o dado certo entregue (testado isolando trecho, testado com
    markdown de tabela também — ambos erraram). Convertendo cada linha em frase própria com o
    valor já rotulado pelo nome da coluna, o modelo passou a ler certo em testes com 2 perguntas
    diferentes, mesmo com cabeçalho imperfeito (coluna mesclada vira "col1", "col2"...). Não é
    caso de "escrever melhor o prompt" — é a mesma informação, só que sem depender do modelo
    reconstruir por conta própria qual número pertence a qual coluna.

    Linha com um só valor preenchido (ex: "CD4051B" sozinho, sub-cabeçalho de seção dentro da
    tabela) vira marcador de seção, não uma "Linha N" — preserva o contexto de qual dispositivo
    aquele bloco de linhas descreve, sem inflar a contagem de linha."""
    dados = tabela.extract()
    if not dados:
        return ""
    cabecalhos = [c if c else f"col{i}" for i, c in enumerate(dados[0])]
    linhas_prosa = []
    numero_linha = 1
    for linha in dados[1:]:
        valores_preenchidos = [(i, v) for i, v in enumerate(linha) if v not in (None, "")]
        if not valores_preenchidos:
            continue
        if len(valores_preenchidos) == 1:
            linhas_prosa.append(f"[Seção: {valores_preenchidos[0][1]}]")
            continue
        pares = [f"{cabecalhos[i]}={v}" for i, v in valores_preenchidos]
        linhas_prosa.append(f"Linha {numero_linha}: " + ", ".join(pares) + ".")
        numero_linha += 1
    return "\n".join(linhas_prosa)


def extrair_texto_pagina_pdf(pagina) -> str:
    """Extração por página: tabela de verdade vira prosa estruturada (`tabela_para_prosa`), resto
    da página continua por `get_text()` normal. Nunca troca a extração inteira por tabela — só
    prosa comum arriscaria falso-positivo de `find_tables()` (ex: lista com marcadores/recuo
    virando "tabela de 1 coluna" por engano) trocar texto corrido por um formato pior pra ele.

    Teto mínimo de 2 linhas de dado × 2 colunas: tabela menor que isso é mais provável falso-
    positivo do que tabela de verdade — nesse caso, ignora a detecção e deixa `get_text()` cuidar
    da região inteira, sem excluir nada."""
    tabelas = [t for t in pagina.find_tables().tables if len(t.extract()) >= 3 and t.col_count >= 2]
    if not tabelas:
        return pagina.get_text()

    bboxes_tabelas = [fitz.Rect(t.bbox) for t in tabelas]
    partes = []
    for bloco in pagina.get_text("blocks"):
        bbox_bloco = fitz.Rect(bloco[:4])
        # pula bloco de texto que cai dentro de alguma tabela detectada — esse conteúdo já vira
        # a versão em prosa da tabela logo abaixo, incluir os dois duplicaria o mesmo dado.
        if any(bbox_bloco.intersects(bbox_tabela) for bbox_tabela in bboxes_tabelas):
            continue
        partes.append(bloco[4])

    for tabela in tabelas:
        prosa = tabela_para_prosa(tabela)
        if prosa:
            partes.append(prosa)

    return "\n".join(partes)


def extrair_texto_pdf(caminho_arquivo: str) -> tuple[str, list[int]]:
    documento = fitz.open(caminho_arquivo)
    partes = []
    offsets_paginas = []
    offset_atual = 0
    for pagina in documento:
        offsets_paginas.append(offset_atual)
        texto_pagina = extrair_texto_pagina_pdf(pagina)
        partes.append(texto_pagina)
        offset_atual += len(texto_pagina) + 1  # +1 pelo "\n" usado no join abaixo
    documento.close()
    return "\n".join(partes), offsets_paginas

def extrair_texto_txt(caminho_arquivo: str) -> tuple[str, list[int]]:
    with open(caminho_arquivo, "r", encoding="utf-8", errors="replace") as f:
        texto = f.read()
    return texto, [0]  # TXT não tem conceito de página

def extrair_texto_docx(caminho_arquivo: str) -> tuple[str, list[int]]:
    documento = docx.Document(caminho_arquivo)
    partes = [paragrafo.text for paragrafo in documento.paragraphs]

    # python-docx separa parágrafos de tabelas — documento.paragraphs sozinho IGNORA
    # tabelas por completo. Documentos técnicos costumam ter dado real em tabela
    # (specs, listas de variáveis), não só texto corrido.
    for tabela in documento.tables:
        for linha in tabela.rows:
            celulas = [celula.text.strip() for celula in linha.cells]
            partes.append(" | ".join(celulas))

    texto = "\n".join(partes)
    return texto, [0]  # Word não pagina sem renderizar — usado só como FALLBACK, ver abaixo


def _localizar_soffice() -> str | None:
    """Acha o executável do LibreOffice — `soffice` no PATH primeiro (Linux/instalação
    customizada), senão o caminho padrão de instalação no Windows. `None` se não achar em
    lugar nenhum (LibreOffice não instalado nesta máquina)."""
    encontrado = shutil.which("soffice") or shutil.which("soffice.exe")
    if encontrado:
        return encontrado
    caminho_padrao_windows = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if os.path.exists(caminho_padrao_windows):
        return caminho_padrao_windows
    return None


def converter_docx_para_pdf(caminho_docx: str) -> str:
    """Converte um .docx pra PDF de verdade via LibreOffice headless (renderiza igual o Word
    faria, calculando quebra de página real) — devolve o caminho do PDF gerado, numa pasta
    temporária que quem chama tem que apagar depois. Levanta `RuntimeError` se não conseguir
    (LibreOffice não instalado, timeout, arquivo corrompido)."""
    soffice = _localizar_soffice()
    if soffice is None:
        raise RuntimeError("LibreOffice (soffice) não encontrado nesta máquina.")

    pasta_saida = tempfile.mkdtemp(prefix="jarvis_docx2pdf_")
    resultado = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", pasta_saida, caminho_docx],
        capture_output=True, text=True, timeout=120,
    )
    nome_pdf = os.path.splitext(os.path.basename(caminho_docx))[0] + ".pdf"
    caminho_pdf = os.path.join(pasta_saida, nome_pdf)
    if resultado.returncode != 0 or not os.path.exists(caminho_pdf):
        raise RuntimeError(f"Conversão DOCX->PDF falhou (código {resultado.returncode}): {resultado.stderr}")
    return caminho_pdf


def extrair_texto_docx_com_paginacao(caminho_arquivo: str) -> tuple[str, list[int]]:
    """DOCX com página REAL — achado real, 2026-08-25: `extrair_texto_docx()` sempre devolvia
    `offsets_paginas=[0]` (página fake, todo chunk caía em "página 1", ver correção
    2026-08-20/2026-08-25 no histórico). Correção de verdade: renderiza o DOCX via LibreOffice
    headless (`converter_docx_para_pdf()`) e reaproveita `extrair_texto_pdf()` — mesmo extrator
    testado e confiável do PDF, incluindo a detecção de tabela (`find_tables()`/`tabela_para_prosa()`).

    Bônus verificado (não só página): testado no `LADDER_MPS_TIA_PORTAL.docx` real — a tabela de
    tags, que saía como "célula1 | célula2" cru pelo `extrair_texto_docx()` antigo, sai como
    "Linha N: campo=valor" (prosa estruturada) pelo caminho novo, porque passa a usar a MESMA
    detecção de tabela já validada pro PDF.

    Cai pro extrator antigo (`extrair_texto_docx`, sem página real) se a conversão falhar —
    LibreOffice pode não estar instalado (outra máquina) ou o arquivo pode dar erro na
    renderização; melhor indexar sem página real do que não indexar nada."""
    try:
        caminho_pdf = converter_docx_para_pdf(caminho_arquivo)
    except (RuntimeError, subprocess.TimeoutExpired) as erro:
        print(f"  [aviso] Conversão DOCX->PDF falhou pra '{caminho_arquivo}' ({erro}) — indexando sem página real.")
        return extrair_texto_docx(caminho_arquivo)

    try:
        return extrair_texto_pdf(caminho_pdf)
    finally:
        pasta_temporaria = os.path.dirname(caminho_pdf)
        os.remove(caminho_pdf)
        os.rmdir(pasta_temporaria)


# TXT sempre devolve offsets_paginas=[0] (todo chunk cai em "página 1") — achado real,
# 2026-08-20 ("transcreva as 3 primeiras páginas" de um .txt silenciosamente devolvia o
# documento inteiro, sem avisar): filtro por página não filtra NADA nesse formato, sem erro nem
# aviso — falha silenciosa. `resumir.py: transcrever_arquivo()` usa esta lista pra recusar o
# filtro por página com uma mensagem honesta, em vez de fingir que funcionou. DOCX tinha o mesmo
# problema até 2026-08-25 — resolvido de verdade com `extrair_texto_docx_com_paginacao()` acima
# (renderiza via LibreOffice, página real), não é mais falha silenciosa, saiu desta lista.
EXTENSOES_SEM_PAGINACAO_REAL = {".txt"}


def possui_estrutura_de_secoes(caminho_arquivo: str, extensao: str) -> bool:
    """Verifica se o documento tem algum sinal de estrutura de seções — sumário embutido no PDF
    (`fitz`/PyMuPDF já expõe isso via `get_toc()`, sem precisar de biblioteca nova) ou parágrafo
    com estilo de título (Heading 1/2/...) no DOCX.

    Achado real, 2026-08-25: testado nos 5 PDFs e 3 DOCX reais do projeto — Os Sertões tem 58
    entradas de sumário (incluindo "Capítulo I" com a página exata) e o datasheet CD405x tem 39;
    Revolução dos Bichos, livro amarelo e steam.pdf não têm nenhuma. Nenhum dos 3 DOCX reais usa
    estilo de Heading. Ou seja: funciona de verdade pra documentos bem estruturados, e
    simplesmente não tem sinal nenhum pros que não são — isso é usado hoje só de forma
    INFORMATIVA (avisar o usuário se faria sentido pedir recorte por capítulo/seção nesse
    documento específico), preparando terreno pra um recorte de verdade por capítulo/seção via
    sumário, ainda não implementado."""
    if extensao == ".pdf":
        documento = fitz.open(caminho_arquivo)
        tem_sumario = len(documento.get_toc()) > 0
        documento.close()
        return tem_sumario
    if extensao == ".docx":
        documento = docx.Document(caminho_arquivo)
        return any(
            p.style.name.startswith("Heading") and p.text.strip()
            for p in documento.paragraphs
        )
    return False


def extrair_secoes_pdf(caminho_arquivo: str) -> list[dict]:
    """Sumário embutido do PDF (`get_toc()`, PyMuPDF) — `[{"titulo", "nivel", "pagina"}, ...]`,
    lista vazia se o PDF não tiver sumário. Mesma fonte de dado que `possui_estrutura_de_secoes()`
    já usa, só devolvendo os detalhes (título/página) em vez de só True/False — usada por
    `resumir.py` pra recortar transcrição por capítulo/seção quando o termo pedido casa com uma
    entrada do sumário."""
    documento = fitz.open(caminho_arquivo)
    toc = documento.get_toc()
    documento.close()
    return [{"titulo": titulo, "nivel": nivel, "pagina": pagina} for nivel, titulo, pagina in toc]


def localizar_titulos_docx(caminho_arquivo: str) -> list[tuple[int, str]]:
    """Parágrafos com estilo de título (Heading 1/2/...) do DOCX, em ordem de leitura —
    `[(nivel, texto), ...]`. Não sabe em que CHUNK cada título cai (depende de como o texto foi
    fatiado na indexação, não é responsabilidade deste extrator) — `resumir.py` localiza a
    posição buscando o texto do título nos chunks já indexados, mesma técnica já usada por
    `termo_busca`."""
    documento = docx.Document(caminho_arquivo)
    titulos = []
    for paragrafo in documento.paragraphs:
        if paragrafo.style.name.startswith("Heading") and paragrafo.text.strip():
            ultima_palavra = paragrafo.style.name.split()[-1]
            nivel = int(ultima_palavra) if ultima_palavra.isdigit() else 1
            titulos.append((nivel, paragrafo.text.strip()))
    return titulos


EXTRATORES_POR_EXTENSAO = {
    ".pdf": extrair_texto_pdf,
    ".txt": extrair_texto_txt,
    ".docx": extrair_texto_docx_com_paginacao,
}

def pagina_do_offset(offset: int, offsets_paginas: list[int]) -> int:
    pagina = 1
    for i, inicio_pagina in enumerate(offsets_paginas):
        if offset >= inicio_pagina:
            pagina = i + 1
        else:
            break
    return pagina

def dividir_em_chunks(texto: str, tamanho: int, sobreposicao: int) -> list[tuple[str, int]]:
    """Retorna lista de (texto_do_chunk, offset_inicial_no_texto_completo)."""
    chunks = []
    inicio = 0
    while inicio < len(texto):
        fim = inicio + tamanho
        if fim < len(texto):
            while fim > inicio and texto[fim] != " ":
                fim -= 1
        chunk = texto[inicio:fim]
        chunks.append((chunk, inicio))
        inicio += tamanho - sobreposicao
    return chunks

def gerar_embeddings_em_lote(textos: list[str]) -> list[list[float]]:
    # MODELO_EMBEDDING vem do config, não escrito à mão: com o nome fixo no código, trocar o
    # modelo no config.py faria documento e histórico serem indexados em espaços vetoriais
    # diferentes — sem erro nenhum, só busca piorando em silêncio.
    resposta = requests.post(
        "http://localhost:11434/api/embed",
        json={"model": MODELO_EMBEDDING, "input": textos}
    )
    return resposta.json()["embeddings"]

def calcular_hash(caminho_arquivo: str) -> str:
    with open(caminho_arquivo, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()

def obter_dados_indexados(nome_arquivo: str, colecao) -> dict | None:
    resultado = colecao.get(where={"arquivo": nome_arquivo}, limit=1)
    if len(resultado["ids"]) == 0:
        return None
    return resultado["metadatas"][0]

def apagar_chunks_do_arquivo(nome_arquivo: str, colecao):
    colecao.delete(where={"arquivo": nome_arquivo})

def indexar_arquivo(nome_arquivo: str, colecao) -> str:
    caminho = os.path.join(PASTA_DOCUMENTOS, nome_arquivo)
    extensao = os.path.splitext(nome_arquivo)[1].lower()
    extrator = EXTRATORES_POR_EXTENSAO.get(extensao)
    if extrator is None:
        return f"'{nome_arquivo}': formato '{extensao}' não suportado pra leitura."

    hash_atual = calcular_hash(caminho)
    dados_indexados = obter_dados_indexados(nome_arquivo, colecao)

    hash_bateu = dados_indexados is not None and dados_indexados.get("hash_conteudo") == hash_atual
    versao_bateu = dados_indexados is not None and dados_indexados.get("pipeline_versao") == VERSAO_PIPELINE

    if hash_bateu and versao_bateu:
        return f"'{nome_arquivo}' sem alterações, pulando."

    if dados_indexados is not None:
        apagar_chunks_do_arquivo(nome_arquivo, colecao)

    texto, offsets_paginas = extrator(caminho)
    chunks_com_offset = dividir_em_chunks(texto, TAMANHO_CHUNK, SOBREPOSICAO)
    chunks = [c for c, _ in chunks_com_offset]

    total_chunks = len(chunks)
    for inicio_lote in range(0, total_chunks, TAMANHO_LOTE_EMBEDDING):
        lote_chunks = chunks[inicio_lote:inicio_lote + TAMANHO_LOTE_EMBEDDING]
        embeddings = gerar_embeddings_em_lote(lote_chunks)

        ids_lote = [f"{nome_arquivo}_{inicio_lote + i}" for i in range(len(lote_chunks))]
        metadatas_lote = [
            {
                "arquivo": nome_arquivo,
                "chunk_num": inicio_lote + i,
                "hash_conteudo": hash_atual,
                "pipeline_versao": VERSAO_PIPELINE,
                "fonte": "documento",
                "pagina": pagina_do_offset(chunks_com_offset[inicio_lote + i][1], offsets_paginas)
            }
            for i in range(len(lote_chunks))
        ]

        colecao.add(
            ids=ids_lote,
            embeddings=embeddings,
            documents=lote_chunks,
            metadatas=metadatas_lote
        )

    return f"'{nome_arquivo}' indexado: {total_chunks} chunks."

if __name__ == "__main__":
    cliente = chromadb.PersistentClient(path=PASTA_BANCO_VETORIAL)
    colecao = cliente.get_or_create_collection(name="documentos_pessoais")

    for nome_arquivo in os.listdir(PASTA_DOCUMENTOS):
        extensao = os.path.splitext(nome_arquivo)[1].lower()
        if extensao not in EXTRATORES_POR_EXTENSAO:
            continue

        print(f"Processando {nome_arquivo}...")
        print(f"  {indexar_arquivo(nome_arquivo, colecao)}")

    print("\nIndexação concluída!")
    print(f"Total de itens no banco: {colecao.count()}")
